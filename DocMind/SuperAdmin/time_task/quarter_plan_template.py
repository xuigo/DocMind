from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from DocMind import db
from DocMind.models import *


def preprocess(quarter):
    results = []
    names = []
    querydata = Qplan.query.filter_by(quarter=quarter).all()
    for data in querydata:
        if data.pname in names:
            continue
        else:
            names.append(data.pname)
    for name in names:
        data_dict = {}
        querydata = Qplan.query.filter_by(quarter=quarter, pname=name).all()
        data_dict['project'] = name
        data_dict['manager'] = querydata[0].manager
        data_dict['quarter'] = quarter
        data_dict['data'] = []
        for data in querydata:
            c_dict = {}
            c_dict['content'] = data.qplanContent
            c_dict['p1'] = data.implementer
            c_dict['p2'] = data.checker1
            c_dict['p3'] = data.checker2
            data_dict['data'].append(c_dict)
        results.append(data_dict)
    return results


def quarter_plan_template(quarter):
    datas = preprocess(quarter)
    document = Document()
    sections = document.sections
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    quarter = quarter
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('威富视界北京研发中心\r {}工作计划'.format(quarter))
    r.font.size = Pt(18)
    r.bold = True
    for data in datas:
        colum = (len(data['data'])+3)*2
        table = document.add_table(rows=colum, cols=17, style='Table Grid')
        table.cell(0, 0).merge(table.cell(1, 16))
        table.cell(2, 0).merge(table.cell(3, 0))
        table.cell(2, 1).merge(table.cell(3, 10))
        table.cell(2, 11).merge(table.cell(3, 12))
        table.cell(2, 13).merge(table.cell(3, 14))
        table.cell(2, 15).merge(table.cell(3, 16))

        hdr_cells0 = table.rows[0].cells
        hdr_cells0[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells0[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells0[0].paragraphs[0].add_run(
            '{}季度计划'.format(data['project']))
        P0.bold = True
        P0.font.size = Pt(14)

        hdr_cells1 = table.rows[2].cells
        hdr_cells1[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1 = hdr_cells1[0].paragraphs[0].add_run('序号')
        P1.bold = True
        P1.font.size = Pt(12)

        hdr_cells1 = table.rows[2].cells
        hdr_cells1[5].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells1[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1 = hdr_cells1[5].paragraphs[0].add_run('工作内容')
        P1.bold = True
        P1.font.size = Pt(12)

        hdr_cells1 = table.rows[2].cells
        hdr_cells1[11].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells1[11].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1 = hdr_cells1[11].paragraphs[0].add_run('实施人')
        P1.bold = True
        P1.font.size = Pt(12)

        hdr_cells1 = table.rows[2].cells
        hdr_cells1[13].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells1[13].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1 = hdr_cells1[13].paragraphs[0].add_run('验收人一')
        P1.bold = True
        P1.font.size = Pt(12)

        hdr_cells1 = table.rows[2].cells
        hdr_cells1[15].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells1[15].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P1 = hdr_cells1[15].paragraphs[0].add_run('验收人二')
        P1.bold = True
        P1.font.size = Pt(12)

        for i in range(len(data['data'])):
            row = i*2+4
            table.cell(row, 0).merge(table.cell(row+1, 0))
            table.cell(row, 1).merge(table.cell(row+1, 10))
            table.cell(row, 11).merge(table.cell(row+1, 12))
            table.cell(row, 13).merge(table.cell(row+1, 14))
            table.cell(row, 15).merge(table.cell(row+1, 16))
            hdr_cells1 = table.rows[row].cells
            hdr_cells1[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            hdr_cells1[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P1 = hdr_cells1[0].paragraphs[0].add_run('{}'.format(i))
            P1.bold = True
            P1.font.size = Pt(12)
            hdr_cells1[5].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            P1 = hdr_cells1[5].paragraphs[0].add_run(
                '{}'.format(data['data'][i]['content']))
            P1.font.size = Pt(12)
            hdr_cells1[11].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            hdr_cells1[11].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P1 = hdr_cells1[11].paragraphs[0].add_run(
                '{}'.format(data['data'][i]['p1']))
            P1.font.size = Pt(12)
            hdr_cells1[13].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            hdr_cells1[13].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P1 = hdr_cells1[13].paragraphs[0].add_run(
                '{}'.format(data['data'][i]['p2']))
            P1.font.size = Pt(12)
            hdr_cells1[15].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            hdr_cells1[15].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            P1 = hdr_cells1[15].paragraphs[0].add_run(
                '{}'.format(data['data'][i]['p3']))
            P1.font.size = Pt(12)
        today = time.localtime(time.time())
        today = time.strftime('%Y-%m-%d', today)
        table.cell(colum-2, 0).merge(table.cell(colum-1, 16))
        hdr_cells0 = table.rows[colum-2].cells
        hdr_cells0[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells0[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells0[0].paragraphs[0].add_run(
            '项目负责人：{}               时间：{}'.format(data['manager'], today))
        P0.font.size = Pt(12)
    filename = 'DocMind/media/docFiles/北京研发中心季度计划_{}.docx'.format(
        time.strftime("%Y-%m-%d", time.localtime()))
    document.save(filename)
    return filename


if __name__ == '__main__':
    quarter_plan_template(quarter)
