from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import datetime
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  

from DocMind.models import *
from DocMind import db


week = '5月份第一周'  # week is a param the value is from admin


def preprocess(time_start, time_end):
    time_start = time_start + ' ' + '0:0:0'
    time_end = time_end + ' ' + '23:59:59'
    time_start = datetime.strptime(time_start, '%Y-%m-%d %H:%M:%S')
    time_end = datetime.strptime(time_end, '%Y-%m-%d %H:%M:%S')
    datas = []
    querydata = Vkreport.query.filter(Vkreport.time >= time_start).filter(
        Vkreport.time <= time_end).all()
    for data in querydata:
        data_dict = {}
        data_ = {}
        data_dict['project'] = data.pname
        data_dict['schedule'] = data.schedule
        data_['last'] = data.lastContent
        data_['this'] = data.thisContent
        data_['next'] = data.nextContent
        data_dict['data'] = data_
        datas.append(data_dict)
    return datas


def weekly_report_template(start_time, end_time):
    datas = preprocess(start_time, end_time)
    document = Document()
    sections = document.sections
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('威富视界北京研发中心\r {}周报汇总'.format(week))
    r.font.size = Pt(18)
    r.bold = True
    colum = 10*len(datas)
    if len(datas) == 0:
        return None
    table = document.add_table(rows=colum, cols=14, style='Table Grid')
    for index, data in enumerate(datas):
        table.cell(0+index*10, 0).merge(table.cell(1+index*10, 1))
        table.cell(0+index*10, 2).merge(table.cell(1+index*10, 13))
        table.cell(2+index*10, 0).merge(table.cell(3+index*10, 1))
        table.cell(2+index*10, 2).merge(table.cell(3+index*10, 13))
        table.cell(4+index*10, 0).merge(table.cell(5+index*10, 1))
        table.cell(4+index*10, 2).merge(table.cell(5+index*10, 13))
        table.cell(6+index*10, 0).merge(table.cell(7+index*10, 1))
        table.cell(6+index*10, 2).merge(table.cell(7+index*10, 13))
        table.cell(8+index*10, 0).merge(table.cell(9+index*10, 1))
        table.cell(8+index*10, 2).merge(table.cell(9+index*10, 8))
        table.cell(8+index*10, 9).merge(table.cell(9+index*10, 13))
        hdr_cells0 = table.rows[0+index*10].cells
        hdr_cells0[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells0[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells0[0].paragraphs[0].add_run('项目名称')
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells0[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells0[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells0[2].paragraphs[0].add_run('{}'.format(data['project']))
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells2 = table.rows[2+index*10].cells
        hdr_cells2[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells2[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells2[0].paragraphs[0].add_run('上周计划')
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells2[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells2[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        P0 = hdr_cells2[2].paragraphs[0].add_run(
            '{}'.format(data['data']['last']))
        hdr_cells4 = table.rows[4+index*10].cells
        hdr_cells4[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells4[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells4[0].paragraphs[0].add_run('本周工作')
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells4[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells4[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        P0 = hdr_cells4[2].paragraphs[0].add_run(
            '{}'.format(data['data']['this']))
        hdr_cells6 = table.rows[6+index*10].cells
        hdr_cells6[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells6[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells6[0].paragraphs[0].add_run('下周计划')
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells6[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells6[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        P0 = hdr_cells6[2].paragraphs[0].add_run(
            '{}'.format(data['data']['next']))
        hdr_cells8 = table.rows[8+index*10].cells
        hdr_cells8[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells8[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells8[0].paragraphs[0].add_run('项目进度')
        P0.bold = True
        P0.font.size = Pt(14)
        hdr_cells8[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells8[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells8[2].paragraphs[0].add_run('{}'.format(data['schedule']))
        today = time.localtime(time.time())
        today = time.strftime('%Y-%m-%d', today)
        hdr_cells8[9].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr_cells8[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P0 = hdr_cells8[9].paragraphs[0].add_run('日期：{}'.format(today))
    filename = 'DocMind/media/docFiles/北京研发中心周报_{}.docx'.format(
        time.strftime("%Y-%m-%d", time.localtime()))
    document.save(filename)
    return filename


if __name__ == '__main__':
    weekly_report_template()
